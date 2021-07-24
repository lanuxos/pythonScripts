# EP.4 Word Automation
import wikipedia
from docx import Document

wikipedia.set_lang('th')
product = ['computer mouse', 'computer keyboard']

document = Document()
document.add_heading('Computer Accessories', 0)

for pd in product:
    content = wikipedia.page(pd).content
    document.add_heading(pd, level=1)
    p = document.add_paragraph('')
    p.add_run(content + '\n\n\n')
    
document.save('../asset/wikipediatoword.docx')
