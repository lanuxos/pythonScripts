from docx import Document
import wikipedia

def Wiki(keyword, lang='th'):
    wikipedia.set_lang(lang)
    data = wikipedia.summary(keyword)
    doc = Document()
    doc.add_heading(keyword, 0)
    doc.add_paragraph(data)
    doc.save(keyword + '.docx')
    print('exported')

Wiki('google','th')
