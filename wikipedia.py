# EP.2 wikipedia gui application
import wikipedia
from tkinter import *
from tkinter import ttk, messagebox
from docx import Document

def Wiki(keyword, lang='th'):
    wikipedia.set_lang(lang)
    data = wikipedia.summary(keyword)
    doc = Document()
    doc.add_heading(keyword, 0)
    doc.add_paragraph(data)
    doc.save(keyword + '_' + lang +'.docx')
    print('exported')

wikipedia.set_lang('th')

GUI = Tk()
GUI.title('Wikipedia Search')
GUI.geometry('500x300')

FONT1 = ('Times New Roman', 15)
FONT2 = ('Courier New', 15)

L = Label(GUI, text='Type the keyword you want to search through Wikipedia').pack()

v_search = StringVar()
E1 = Entry(GUI, textvariable=v_search, width=50, justify='center')
E1.pack(ipady=10)

def Search(event=None):
    keyword = v_search.get()
    try:
        language = v_radio.get()
        Wiki(keyword, language)
        messagebox.showinfo('Exported', 'The search result had been save to docx file.')
    except:
        messagebox.showwarning('Keyword Error', 'Please try another keyword search')
    result = wikipedia.summary(keyword)
    print(result)

GUI.bind('<Return>', Search)

F1 = Frame(GUI)
F1.pack(ipadx=10, ipady=10)

v_radio = StringVar()
RB1 = Radiobutton(F1, text='Thai', variable=v_radio, value='th')
RB2 = Radiobutton(F1, text='Chinese', variable=v_radio, value='zh')
RB3 = Radiobutton(F1, text='English', variable=v_radio, value='en')

RB1.invoke()

RB1.grid(row=0, column=0)
RB2.grid(row=0, column=1)
RB3.grid(row=0, column=2)

B1 = ttk.Button(GUI, text='Search', command=Search)
B1.pack(ipadx=20, ipady=10)

B2 = ttk.Button(GUI, text='Exit', command=GUI.destroy).pack(ipadx=10, ipady=10)

GUI.mainloop()
